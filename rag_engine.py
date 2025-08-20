from llama_index.core import VectorStoreIndex, Document, StorageContext, load_index_from_storage
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core import Settings
import os

# ----------------- Global Config -----------------
PERSIST_DIR = "vector_db"

# Embedding + LLM
Settings.embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
llm = Ollama(model="llama3", request_timeout=120.0)   # ✅ use llama3 locally

# ----------------- File Ingestion -----------------
def ingest_file(uploaded_file):
    from PyPDF2 import PdfReader
    import docx

    if uploaded_file.type == "application/pdf":
        pdf = PdfReader(uploaded_file)
        text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        raise Exception("Unsupported file type")

    # Split into chunks
    parser = SimpleNodeParser.from_defaults(chunk_size=512, chunk_overlap=50)
    nodes = parser.get_nodes_from_documents([Document(text=text)])

    # Save vector DB
    index = VectorStoreIndex(nodes)
    index.storage_context.persist(persist_dir=PERSIST_DIR)
    return len(nodes)

# ----------------- QA Function -----------------
def answer_question(query):
    if not os.path.exists(PERSIST_DIR):
        return "❌ No document found. Please upload a file first.", []

    # Load stored index
    storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
    index = load_index_from_storage(storage_context)

    # Retrieve chunks
    retriever = index.as_retriever(similarity_top_k=3)
    retrieved_nodes = retriever.retrieve(query)

    if not retrieved_nodes:
        return "No relevant information found in the uploaded document.", []

    # Combine retrieved context
    context = "\n\n".join([n.node.get_content() for n in retrieved_nodes])

    # ✅ Ask LLM to explain based only on retrieved context
    prompt = f"""
You are an assistant that answers questions ONLY from the provided context. 
Do not add information that is not in the context.  
Explain the answer clearly in simple words.  

Question: {query}

Context from uploaded document:
{context}

Answer (with explanation):
"""
    response = llm.complete(prompt)

    return response.text.strip(), [n.node.get_content() for n in retrieved_nodes]
